# -*- coding: utf-8 -*-
"""Build a Hebrew RTL .docx from the PDF article."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"C:\Users\USER\Desktop\עוסק-פטור-או-עוסק-מורשה-v2.docx"

NAVY = RGBColor(0x1E, 0x3A, 0x5F)
TEAL = RGBColor(0x14, 0xB8, 0xA6)
DARK = RGBColor(0x1F, 0x29, 0x37)
MUTED = RGBColor(0x6B, 0x72, 0x80)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
FONT = 'David'


def _el(tag, **attrs):
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn(k), v)
    return e


def set_doc_defaults_rtl(doc):
    """Force RTL + Hebrew font at the document-default level."""
    styles_el = doc.styles.element
    docDefaults = styles_el.find(qn('w:docDefaults'))
    if docDefaults is None:
        docDefaults = _el('w:docDefaults')
        styles_el.insert(0, docDefaults)

    # Remove existing defaults
    for tag in ('w:rPrDefault', 'w:pPrDefault'):
        old = docDefaults.find(qn(tag))
        if old is not None:
            docDefaults.remove(old)

    rPrDefault = _el('w:rPrDefault')
    rPr = _el('w:rPr')
    rFonts = _el('w:rFonts')
    rFonts.set(qn('w:ascii'), FONT)
    rFonts.set(qn('w:hAnsi'), FONT)
    rFonts.set(qn('w:cs'), FONT)
    rFonts.set(qn('w:hint'), 'cs')
    rPr.append(rFonts)
    rPr.append(_el('w:rtl'))
    sz = _el('w:sz'); sz.set(qn('w:val'), '22'); rPr.append(sz)
    szCs = _el('w:szCs'); szCs.set(qn('w:val'), '22'); rPr.append(szCs)
    lang = _el('w:lang')
    lang.set(qn('w:val'), 'he-IL')
    lang.set(qn('w:bidi'), 'he-IL')
    rPr.append(lang)
    rPrDefault.append(rPr)

    pPrDefault = _el('w:pPrDefault')
    pPr = _el('w:pPr')
    pPr.append(_el('w:bidi'))
    jc = _el('w:jc'); jc.set(qn('w:val'), 'right'); pPr.append(jc)
    pPrDefault.append(pPr)

    docDefaults.append(rPrDefault)
    docDefaults.append(pPrDefault)


def set_section_rtl(section):
    sectPr = section._sectPr
    bidi = sectPr.find(qn('w:bidi'))
    if bidi is None:
        sectPr.append(_el('w:bidi'))
    rtlGutter = sectPr.find(qn('w:rtlGutter'))
    if rtlGutter is None:
        sectPr.append(_el('w:rtlGutter'))


def set_para_rtl(paragraph, align='right'):
    pPr = paragraph._p.get_or_add_pPr()
    # remove any existing bidi/jc
    for tag in ('w:bidi', 'w:jc'):
        old = pPr.find(qn(tag))
        if old is not None:
            pPr.remove(old)
    pPr.append(_el('w:bidi'))
    jc = _el('w:jc'); jc.set(qn('w:val'), align); pPr.append(jc)
    if align == 'right':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == 'center':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def shade(element_tcPr_or_pPr, color_hex):
    shd = _el('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    element_tcPr_or_pPr.append(shd)


def shade_cell(cell, color_hex):
    shade(cell._tc.get_or_add_tcPr(), color_hex)


def shade_paragraph(paragraph, color_hex):
    shade(paragraph._p.get_or_add_pPr(), color_hex)


def set_table_rtl(table):
    tblPr = table._element.tblPr
    if tblPr.find(qn('w:bidiVisual')) is None:
        tblPr.append(_el('w:bidiVisual'))


def add_run(paragraph, text, bold=False, color=None, size=11, italic=False):
    run = paragraph.add_run(text)
    rPr = run._element.get_or_add_rPr()
    # Hebrew font on all scripts
    for old in rPr.findall(qn('w:rFonts')):
        rPr.remove(old)
    rFonts = _el('w:rFonts')
    rFonts.set(qn('w:ascii'), FONT)
    rFonts.set(qn('w:hAnsi'), FONT)
    rFonts.set(qn('w:cs'), FONT)
    rFonts.set(qn('w:hint'), 'cs')
    rPr.append(rFonts)
    # size (both Latin + complex-script)
    sz = _el('w:sz'); sz.set(qn('w:val'), str(size * 2)); rPr.append(sz)
    szCs = _el('w:szCs'); szCs.set(qn('w:val'), str(size * 2)); rPr.append(szCs)
    if bold:
        rPr.append(_el('w:b'))
        rPr.append(_el('w:bCs'))
    if italic:
        rPr.append(_el('w:i'))
        rPr.append(_el('w:iCs'))
    rPr.append(_el('w:rtl'))
    lang = _el('w:lang')
    lang.set(qn('w:val'), 'he-IL')
    lang.set(qn('w:bidi'), 'he-IL')
    rPr.append(lang)
    if color:
        run.font.color.rgb = color
    return run


def add_para(doc, text='', size=11, bold=False, color=None, space_after=6,
             align='right', line_spacing=1.5, space_before=0):
    p = doc.add_paragraph()
    set_para_rtl(p, align=align)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = line_spacing
    if text:
        add_run(p, text, bold=bold, size=size, color=color)
    return p


def add_heading(doc, text, level=2):
    """Heading navy bold with a short teal underline bar (matches PDF)."""
    sizes = {1: 26, 2: 20, 3: 14}
    size = sizes.get(level, 14)
    # Title paragraph
    p = doc.add_paragraph()
    set_para_rtl(p, align='right')
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, text, bold=True, size=size, color=NAVY)
    # Short teal bar beneath via separate paragraph with bottom border and
    # right indent so the line is short (~5cm from the right)
    bar = doc.add_paragraph()
    set_para_rtl(bar, align='right')
    bar.paragraph_format.space_before = Pt(0)
    bar.paragraph_format.space_after = Pt(10)
    bar.paragraph_format.line_spacing = 1.0
    bar.paragraph_format.left_indent = Cm(11)  # leave ~5cm bar on the right side
    pPr = bar._p.get_or_add_pPr()
    pBdr = _el('w:pBdr')
    bottom = _el('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '18')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '14B8A6')
    pBdr.append(bottom)
    pPr.append(pBdr)
    # empty run so paragraph renders
    r = bar.add_run('')
    r.font.size = Pt(1)
    return p


def add_bullet(doc, parts):
    p = doc.add_paragraph(style='List Bullet')
    set_para_rtl(p, align='right')
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(3)
    for text, bold in parts:
        add_run(p, text, bold=bold, size=11)
    return p


def add_numbered(doc, parts, num):
    p = doc.add_paragraph()
    set_para_rtl(p, align='right')
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    add_run(p, f"{num}. ", bold=True, size=11, color=NAVY)
    for text, bold in parts:
        add_run(p, text, bold=bold, size=11)


def add_callout(doc, icon, title, body_parts, bg='E6F7F4', border='14B8A6', text_color=None):
    """Callout with soft teal bg + thick teal bar on the START side (right in RTL)."""
    t = doc.add_table(rows=1, cols=1)
    set_table_rtl(t)
    # pad the cell
    cell = t.cell(0, 0)
    shade_cell(cell, bg)
    tcPr = cell._tc.get_or_add_tcPr()
    # cell margins
    tcMar = _el('w:tcMar')
    for edge, val in (('top', '140'), ('bottom', '140'), ('left', '200'), ('right', '200')):
        m = _el(f'w:{edge}')
        m.set(qn('w:w'), val)
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    tcPr.append(tcMar)
    # borders – thick teal only on start (RTL start = right in cell XML sense? Word uses start/end)
    tcBorders = _el('w:tcBorders')
    for edge, sz, color in (
        ('top', '4', 'B2E5DB'),
        ('bottom', '4', 'B2E5DB'),
        ('left', '4', 'B2E5DB'),   # logical end-side faint
        ('right', '36', '14B8A6'), # thick teal bar visible on right (RTL leading edge)
    ):
        b = _el(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), sz)
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)
    # content
    p = cell.paragraphs[0]
    set_para_rtl(p, align='right')
    p.paragraph_format.line_spacing = 1.5
    add_run(p, f"{icon} ", bold=True, size=11)
    add_run(p, title, bold=True, size=11, color=NAVY if not text_color else text_color)
    add_run(p, " ", size=11)
    for text, bold in body_parts:
        add_run(p, text, bold=bold, size=11, color=text_color)
    add_para(doc, space_after=8)
    return t


def _cell_borders(cell, top_sz='0', bottom_sz='4', side='nil',
                  top_color='E5E7EB', bottom_color='E5E7EB'):
    tcPr = cell._tc.get_or_add_tcPr()
    # remove existing borders
    for old in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(old)
    tcBorders = _el('w:tcBorders')
    defs = [
        ('top', top_sz, top_color),
        ('bottom', bottom_sz, bottom_color),
        ('left', '0', 'auto'),
        ('right', '0', 'auto'),
    ]
    for edge, sz, color in defs:
        b = _el(f'w:{edge}')
        if sz == '0' or side == 'nil':
            if edge in ('left', 'right'):
                b.set(qn('w:val'), 'nil')
                tcBorders.append(b)
                continue
        if sz == '0':
            b.set(qn('w:val'), 'nil')
        else:
            b.set(qn('w:val'), 'single')
            b.set(qn('w:sz'), sz)
            b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcMar')):
        tcPr.remove(old)
    tcMar = _el('w:tcMar')
    for edge, val in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
        m = _el(f'w:{edge}')
        m.set(qn('w:w'), str(val))
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    tcPr.append(tcMar)


def _cell_vcenter(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:vAlign')):
        tcPr.remove(old)
    v = _el('w:vAlign'); v.set(qn('w:val'), 'center'); tcPr.append(v)


def _set_table_no_outer_borders(table):
    """Minimal clean look – only horizontal row separators."""
    tblPr = table._element.tblPr
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)
    tblBorders = _el('w:tblBorders')
    for edge, val in (('top', 'nil'), ('bottom', 'nil'), ('left', 'nil'),
                      ('right', 'nil'), ('insideV', 'nil')):
        b = _el(f'w:{edge}')
        b.set(qn('w:val'), val)
        tblBorders.append(b)
    inH = _el('w:insideH')
    inH.set(qn('w:val'), 'single')
    inH.set(qn('w:sz'), '4')
    inH.set(qn('w:color'), 'E5E7EB')
    tblBorders.append(inH)
    tblPr.append(tblBorders)


def add_comparison_table(doc, rows):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    set_table_rtl(t)
    _set_table_no_outer_borders(t)
    # full width
    tblW = _el('w:tblW')
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')
    t._element.tblPr.append(tblW)

    # Header row
    header_cells = t.rows[0].cells
    for i, h in enumerate(rows[0]):
        cell = header_cells[i]
        shade_cell(cell, '1E3A5F')
        _cell_borders(cell, top_sz='0', bottom_sz='0')
        _cell_margins(cell, top=140, bottom=140)
        _cell_vcenter(cell)
        p = cell.paragraphs[0]
        set_para_rtl(p, align='center' if i > 0 else 'right')
        p.paragraph_format.line_spacing = 1.15
        add_run(p, h, bold=True, size=11, color=WHITE)

    # Body rows
    for ri, row in enumerate(rows[1:], start=1):
        for ci, v in enumerate(row):
            cell = t.rows[ri].cells[ci]
            _cell_borders(cell)
            _cell_margins(cell, top=120, bottom=120)
            _cell_vcenter(cell)
            p = cell.paragraphs[0]
            set_para_rtl(p, align='right' if ci == 0 else 'center')
            p.paragraph_format.line_spacing = 1.2
            is_first = (ci == 0)
            add_run(p, v, bold=is_first, size=10,
                    color=NAVY if is_first else DARK)
    return t


def build():
    doc = Document()
    set_doc_defaults_rtl(doc)

    # Normal style tweak
    normal = doc.styles['Normal']
    normal.font.name = FONT
    normal.font.size = Pt(11)

    for s in doc.sections:
        s.top_margin = Cm(2)
        s.bottom_margin = Cm(2)
        s.left_margin = Cm(2.2)
        s.right_margin = Cm(2.2)
        set_section_rtl(s)

    # Badge – small teal pill in a 1-cell table, aligned right (top of page)
    badge_t = doc.add_table(rows=1, cols=1)
    set_table_rtl(badge_t)
    # narrow width
    tblW = _el('w:tblW'); tblW.set(qn('w:w'), '2600'); tblW.set(qn('w:type'), 'dxa')
    badge_t._element.tblPr.append(tblW)
    # right-align table inside page
    jc = _el('w:jc'); jc.set(qn('w:val'), 'right')
    badge_t._element.tblPr.append(jc)
    # remove borders
    tblBorders = _el('w:tblBorders')
    for e in ('top', 'bottom', 'left', 'right', 'insideH', 'insideV'):
        b = _el(f'w:{e}'); b.set(qn('w:val'), 'nil'); tblBorders.append(b)
    badge_t._element.tblPr.append(tblBorders)

    bcell = badge_t.cell(0, 0)
    shade_cell(bcell, '14B8A6')
    _cell_margins(bcell, top=80, bottom=80, left=200, right=200)
    bp = bcell.paragraphs[0]
    set_para_rtl(bp, align='center')
    bp.paragraph_format.space_after = Pt(0)
    add_run(bp, "מדריך מעודכן 2026", size=10, bold=True, color=WHITE)
    add_para(doc, space_after=8)

    # Title
    p = doc.add_paragraph()
    set_para_rtl(p, align='right')
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    add_run(p, "עוסק פטור או עוסק מורשה – מה ההבדל ואיזה סוג עסק כדאי לפתוח?",
            bold=True, size=26, color=NAVY)

    add_para(doc, "מאת: צוות פרפקט וואן | עודכן: אפריל 2026 | זמן קריאה: 8 דקות",
             size=10, color=MUTED, space_after=18)

    # Intro
    p = add_para(doc, space_after=10)
    add_run(p, "פתיחת עסק בישראל מחייבת החלטה חשובה כבר בתחילת הדרך: ")
    add_run(p, "האם לפתוח עוסק פטור או עוסק מורשה?", bold=True)

    p = add_para(doc, space_after=10)
    add_run(p,
        "הבחירה בין שני סוגי העסקים משפיעה על תשלום המסים, גביית מע\"מ, ניהול החשבונות "
        "והיקף הפעילות העסקית. טעות בשלב הזה יכולה לעלות אלפי שקלים בשנה — ולכן חשוב "
        "להבין את ההבדלים ")
    add_run(p, "לפני", bold=True)
    add_run(p, " שפותחים תיק.")

    add_para(doc,
        "במאמר הזה נסביר בצורה ברורה ומעשית את ההבדלים בין עוסק פטור לעוסק מורשה, למי מתאים "
        "כל סוג עסק, כמה מס משלמים בכל מסלול, ומה צריך לדעת לפני שמתחילים.",
        space_after=16)

    # Section: עוסק פטור
    add_heading(doc, "מה זה עוסק פטור?")
    p = add_para(doc, space_after=10)
    add_run(p, "עוסק פטור הוא סוג של עסק קטן שמיועד לעצמאים עם ")
    add_run(p, "מחזור הכנסות שנתי של עד 120,000 ₪", bold=True)
    add_run(p, " (נכון ל-2026).")

    p = add_para(doc, space_after=6)
    add_run(p, "היתרון המרכזי של עוסק פטור הוא ")
    add_run(p, "הפשטות", bold=True)
    add_run(p, ":")

    add_bullet(doc, [("העסק ", False), ("אינו גובה מע\"מ", True), (" מהלקוחות", False)])
    add_bullet(doc, [("אין חובת הגשת דוח מע\"מ דו-חודשי", False)])
    add_bullet(doc, [("ניהול ספרים פשוט — ספר הכנסות בסיסי", False)])
    add_bullet(doc, [("הדיווח למע\"מ הוא ", False), ("פעם בשנה בלבד", True)])
    add_bullet(doc, [("אין צורך ברואה חשבון (אם כי מומלץ)", False)])

    p = add_para(doc, space_after=10)
    add_run(p, "מצד שני, עוסק פטור ")
    add_run(p, "לא יכול לקזז מע\"מ", bold=True)
    add_run(p, " על הוצאות עסקיות, ולקוחות עסקיים לא יכולים לקזז מע\"מ על התשלומים לו.")

    p = add_para(doc, space_after=10)
    add_run(p, "מי שרוצה להבין את התהליך המלא יכול לקרוא מדריך מפורט על ")
    add_run(p, "איך פותחים עוסק פטור — שלב אחר שלב", bold=True, color=TEAL)
    add_run(p, " באתר פרפקט וואן.")

    # Section: עוסק מורשה
    add_heading(doc, "מה זה עוסק מורשה?")
    p = add_para(doc, space_after=10)
    add_run(p, "עוסק מורשה הוא סוג עסק שבו בעל העסק ")
    add_run(p, "גובה 18% מע\"מ מהלקוחות ומעביר אותו לרשויות המס", bold=True)
    add_run(p, ". סוג זה מתאים לעסקים עם מחזור הכנסות גבוה יותר או לעסקים שעובדים מול חברות וארגונים.")

    add_para(doc, "היתרונות המרכזיים של עוסק מורשה:", space_after=6)
    add_bullet(doc, [("אין תקרת הכנסות", True), (" — אפשר להרוויח ללא הגבלה", False)])
    add_bullet(doc, [("קיזוז מע\"מ", True), (" על כל הוצאה עסקית (ציוד, שכירות, דלק, תוכנות)", False)])
    add_bullet(doc, [("לקוחות עסקיים ", False), ("מעדיפים לעבוד עם מורשה", True),
                     (" כי הם יכולים לקזז את המע\"מ", False)])
    add_bullet(doc, [("מנפיק ", False), ("חשבונית מס", True), (" (לא חשבונית עסקה)", False)])
    add_bullet(doc, [("נתפס ", False), ("כמקצועי יותר", True), (" בעיני לקוחות וספקים", False)])

    p = add_para(doc, space_after=10)
    add_run(p, "החיסרון: ")
    add_run(p, "ניהול מורכב יותר", bold=True)
    add_run(p, " — דוח מע\"מ דו-חודשי, הנהלת חשבונות מסודרת, ולרוב צריך רואה חשבון.")

    p = add_para(doc, space_after=10)
    add_run(p, "מי שמתכנן פעילות עסקית רחבה יותר יכול לבדוק את ")
    add_run(p, "המדריך המלא לפתיחת עוסק מורשה", bold=True, color=TEAL)
    add_run(p, " באתר פרפקט וואן.")

    # Comparison table
    add_heading(doc, "טבלת השוואה — עוסק פטור מול עוסק מורשה")
    add_comparison_table(doc, [
        ("נושא", "עוסק פטור", "עוסק מורשה"),
        ("תקרת הכנסה שנתית", "120,000 ₪ (2026)", "ללא הגבלה"),
        ("גביית מע\"מ", "לא גובה", "גובה 18% מע\"מ"),
        ("קיזוז מע\"מ על הוצאות", "לא יכול", "מקזז מע\"מ תשומות"),
        ("דיווח למע\"מ", "פעם בשנה", "כל חודשיים"),
        ("סוג חשבונית", "חשבונית עסקה / קבלה", "חשבונית מס"),
        ("ניהול ספרים", "ספר הכנסות בלבד", "הנהלת חשבונות מלאה"),
        ("צורך ברואה חשבון", "מומלץ, לא חובה", "מומלץ מאוד"),
        ("עלות רו\"ח חודשית", "75–150 ₪", "200–800 ₪"),
        ("עבודה מול חברות", "מוגבלת (לקוח לא מקזז מע\"מ)", "מועדפת"),
        ("מתאים ל...", "עסק קטן, הכנסה צדדית, התחלה", "עסק צומח, פרילנס מלא, B2B"),
    ])
    add_para(doc, space_after=12)

    # Numeric
    add_heading(doc, "כמה מס משלמים? השוואה מספרית")
    p = add_para(doc, space_after=10)
    add_run(p, "הנה השוואה מעשית — ")
    add_run(p, "כמה נשאר נטו", bold=True)
    add_run(p, " בכל מסלול:")

    add_comparison_table(doc, [
        ("הכנסה חודשית", "עוסק פטור — נטו", "עוסק מורשה — נטו", "הערה"),
        ("5,000 ₪", "~4,700 ₪", "~4,500 ₪", "פטור עדיף"),
        ("8,000 ₪", "~7,200 ₪", "~7,000 ₪", "פטור עדיף"),
        ("10,000 ₪", "~8,500 ₪", "~8,600 ₪", "כמעט שווה"),
        ("12,000 ₪", "חורג מתקרה!", "~10,000 ₪", "חייב מורשה"),
        ("20,000 ₪", "—", "~14,800 ₪", "רק מורשה"),
    ])
    add_para(doc, space_after=12)

    add_callout(doc, "💡", "נקודת מפתח:",
                [("אם ההכנסות שלכם מתקרבות ל-10,000 ₪ בחודש ויש לכם הוצאות עסקיות משמעותיות "
                  "(ציוד, תוכנות, נסיעות) — ", False),
                 ("עוסק מורשה עשוי להיות משתלם יותר", True),
                 (" בזכות קיזוז המע\"מ.", False)])

    # למי מתאים עוסק פטור
    add_heading(doc, "למי מתאים עוסק פטור?")
    add_para(doc, "עוסק פטור מתאים בעיקר למי שמתחיל פעילות עסקית קטנה:", space_after=6)
    add_bullet(doc, [("פרילנסרים בתחילת הדרך", True), (" — כותבים, מעצבים, מתרגמים", False)])
    add_bullet(doc, [("מורים פרטיים", True), (" — שיעורים פרטיים, חוגים", False)])
    add_bullet(doc, [("נותני שירותים קטנים", True), (" — הדברה, תיקונים, שליחויות", False)])
    add_bullet(doc, [("שכירים עם הכנסה נוספת", True), (" — עבודה צדדית בנוסף למשרה", False)])
    add_bullet(doc, [("מוכרים אונליין", True), (" — eBay, Etsy, מכירות קטנות ברשת", False)])
    add_para(doc,
        "מחפשים מדריך ספציפי למקצוע שלכם? יש מדריכים ייעודיים לשליחי וולט, מעצבים גרפיים, "
        "מורים פרטיים ועוד.",
        space_after=12)

    # למי מתאים עוסק מורשה
    add_heading(doc, "למי מתאים עוסק מורשה?")
    add_para(doc, "עוסק מורשה מתאים לעסקים עם פעילות רחבה יותר:", space_after=6)
    add_bullet(doc, [("פרילנסרים שמרוויחים מעל 10,000 ₪/חודש", True)])
    add_bullet(doc, [("עסקים שעובדים מול חברות", True), (" — הלקוחות דורשים חשבונית מס", False)])
    add_bullet(doc, [("בעלי מקצועות חופשיים", True),
                     (" — עורכי דין, רואי חשבון, רופאים (חייבים להיות מורשים)", False)])
    add_bullet(doc, [("יבואנים וסוחרים", True), (" — חייבים מורשה לפי חוק", False)])
    add_bullet(doc, [("עסקים עם הוצאות גבוהות", True),
                     (" — שכירות, ציוד, רכב (קיזוז מע\"מ משתלם)", False)])
    add_para(doc, space_after=8)

    # 5 שאלות
    add_heading(doc, "5 שאלות שיעזרו לכם להחליט")
    questions = [
        [("כמה אני צפוי להרוויח בשנה?", True), (" — מעל 120,000 ₪? חייב מורשה.", False)],
        [("הלקוחות שלי עסקים או פרטיים?", True), (" — עסקים דורשים חשבונית מס = מורשה.", False)],
        [("יש לי הוצאות עסקיות משמעותיות?", True),
         (" — הוצאות גבוהות + קיזוז מע\"מ = מורשה עדיף.", False)],
        [("אני רוצה פשטות מקסימלית?", True), (" — כן = פטור.", False)],
        [("אני מתכנן לגדול מהר?", True), (" — כן = עדיף להתחיל מורשה ולחסוך מעבר.", False)],
    ]
    for i, q in enumerate(questions, start=1):
        add_numbered(doc, q, i)
    add_para(doc, space_after=6)

    add_callout(doc, "⚡", "טיפ חשוב:",
                [("תמיד אפשר ", False),
                 ("להתחיל כעוסק פטור ולעבור למורשה", True),
                 (" בהמשך. המעבר פשוט ואפשר לבצע אותו תוך ימים. ", False),
                 ("אבל לעבור מעוסק מורשה חזרה לפטור — זה הרבה יותר מסובך.", True)])

    # How to open
    add_heading(doc, "איך לפתוח עוסק פטור או מורשה?")
    add_para(doc, "פתיחת עוסק פטור", size=12, bold=True, color=NAVY, space_after=4)
    add_para(doc,
        "התהליך פשוט וניתן לביצוע אונליין תוך 1-3 ימי עסקים — כולל מסמכים נדרשים, טפסים ועלויות.",
        space_after=10)
    add_para(doc, "פתיחת עוסק מורשה", size=12, bold=True, color=NAVY, space_after=4)
    add_para(doc,
        "דורש רישום בשלושה גופים: מע\"מ, מס הכנסה, וביטוח לאומי. התהליך לוקח 3-10 ימי עסקים "
        "וכולל 7 שלבי רישום.",
        space_after=14)

    # CTA dark – full navy block, centered title + subtitle + teal phone
    t = doc.add_table(rows=1, cols=1)
    set_table_rtl(t)
    tblW = _el('w:tblW'); tblW.set(qn('w:w'), '5000'); tblW.set(qn('w:type'), 'pct')
    t._element.tblPr.append(tblW)
    # remove table borders
    tblBorders = _el('w:tblBorders')
    for e in ('top', 'bottom', 'left', 'right', 'insideH', 'insideV'):
        b = _el(f'w:{e}'); b.set(qn('w:val'), 'nil'); tblBorders.append(b)
    t._element.tblPr.append(tblBorders)

    cell = t.cell(0, 0)
    shade_cell(cell, '1E3A5F')
    _cell_margins(cell, top=300, bottom=300, left=300, right=300)
    _cell_vcenter(cell)

    p = cell.paragraphs[0]
    set_para_rtl(p, align='center')
    p.paragraph_format.space_after = Pt(6)
    add_run(p, "רוצים עזרה בבחירה?", bold=True, size=16, color=WHITE)

    p2 = cell.add_paragraph()
    set_para_rtl(p2, align='center')
    p2.paragraph_format.space_after = Pt(8)
    p2.paragraph_format.line_spacing = 1.4
    add_run(p2,
        "צוות רואי החשבון שלנו בפרפקט וואן יעזור לכם לבחור את המסלול הנכון — בחינם וללא התחייבות.",
        size=11, color=WHITE)

    p3 = cell.add_paragraph()
    set_para_rtl(p3, align='center')
    p3.paragraph_format.space_after = Pt(0)
    add_run(p3, "לפרטים: 050-227-7087", bold=True, size=14, color=TEAL)
    add_para(doc, space_after=14)

    # Summary
    add_heading(doc, "לסיכום")
    p = add_para(doc, space_after=10)
    add_run(p, "עוסק פטור ועוסק מורשה הם שני מסלולים שונים לפתיחת עסק בישראל. ")
    add_run(p, "אין מסלול \"טוב\" או \"רע\"", bold=True)
    add_run(p, " — יש מסלול שמתאים לך.")

    add_para(doc, "כלל אצבע:", size=11, bold=True, space_after=4)
    add_bullet(doc, [("הכנסה נמוכה + לקוחות פרטיים + רוצה פשטות = ", False), ("עוסק פטור", True)])
    add_bullet(doc, [("הכנסה גבוהה + לקוחות עסקיים + הוצאות + צמיחה = ", False), ("עוסק מורשה", True)])

    p = add_para(doc, space_after=10)
    add_run(p, "לפני שמקימים עסק חדש, כדאי להבין היטב את היתרונות והחסרונות של כל מסלול. וזכרו — ")
    add_run(p, "תמיד אפשר לעבור ממסלול למסלול", bold=True)
    add_run(p, ".")

    add_para(doc, "למידע נוסף ומדריכים מקיפים, בקרו באתר פרפקט וואן — perfect1.co.il.",
             space_after=16)

    # Separator
    p = add_para(doc, space_after=10)
    pPr = p._p.get_or_add_pPr()
    pBdr = _el('w:pBdr')
    top = _el('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '6')
    top.set(qn('w:space'), '1')
    top.set(qn('w:color'), 'D1D5DB')
    pBdr.append(top)
    pPr.append(pBdr)

    # Footer
    p = add_para(doc, space_after=4, align='center')
    add_run(p, "מאמר זה נכתב על ידי צוות ", size=10, color=MUTED)
    add_run(p, "פרפקט וואן", size=10, bold=True, color=NAVY)
    add_run(p, " — פורטל מידע מקצועי לעצמאים ובעלי עסקים בישראל", size=10, color=MUTED)

    p = add_para(doc, space_after=4, align='center')
    add_run(p, "www.perfect1.co.il | 050-227-7087", size=10, color=MUTED)

    doc.save(OUT)
    print(f"saved: {OUT}")


if __name__ == "__main__":
    build()
